# Test program created while learning the uses of the modules
import docx,docs
d = docx.Document('demo.docx')
print(d.paragraphs)
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
p=d.paragraphs[1]
print(len(p.runs))
print('-------------')
for x in range(len(p.runs)):
    if(p.runs[x].bold == True):
        print(p.runs[x].text)
print('-------------')
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[1].bold)
print(p.runs[3].italic)
print(p.runs[0].bold == None)
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'
p.style = 'Title'
d.save('demo2.docx')
