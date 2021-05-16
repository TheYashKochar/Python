# Program to Create a new Word File
import docx,docs
d = docx.Document()
d.add_paragraph('Random sentence lorem ipsum dolor sit idk.')
d.add_paragraph('Random sentence lorem ipsum dolor sit idk.')
d.add_paragraph('Random sentence lorem ipsum dolor sit idk.')
p = d.paragraphs[0]
p.add_run('This is a new run')
p.runs[1].bold = True
d.save('New Document.docx')
